

from PyQt5.QtGui import *
from PyQt5.QtWidgets import *
from PyQt5.QtCore import *
from PyQt5 import QtCore, QtGui, QtWidgets
from PyQt5.QtWidgets import QMessageBox
from PyQt5.QtGui import QMovie

from docx import Document
from docx.shared import Inches

import smtplib 
from email.mime.multipart import MIMEMultipart 
from email.mime.text import MIMEText 
from email.mime.base import MIMEBase 
from email import encoders

import tensorflow as tf
import sys
import os
import suggest

# Disable tensorflow compilation warnings
os.environ['TF_CPP_MIN_LOG_LEVEL']='2'
import tensorflow as tf

import tensorflow.compat.v1 as tf
tf.disable_v2_behavior()

class Ui_Analyse(object):
    home=0
    def __init__(self,patient,age,number,email):
        self.patient=patient
        self.age=age
        self.number=number
        self.mailid=email
    def setupUi(self, Analyse):
        self.experts={'braintumor':["Dr. Shibam","mallick95shibam@gmail.com"],'stomachcancer':["Dr.Rajendra","rajendra20023@gmail.com"],'lungcancer':["Dr. Ranjan",'ranjankc148@gmail.com']}
        self.imagestatus="Select Scanned Image"
        Analyse.setObjectName("Analyse")
        Analyse.resize(580, 640)
        Analyse.setFixedSize(QtCore.QSize(580,640))
        self.result=''
        self.mailnotify=''
        self.buttonText='Study Image'
        self.label = QtWidgets.QLabel(Analyse)
        self.label.setGeometry(QtCore.QRect(0, 0, 580, 640))
        self.label.setText("")
        self.label.setScaledContents(True)
        self.label.setObjectName("label")
        self.gif2=QMovie("Gui-Images/doctor_checkup.gif")
        self.label.setMovie(self.gif2)
        self.gif2.start()
        self.label_2 = QtWidgets.QLabel(Analyse)
        self.label_2.setGeometry(QtCore.QRect(80, 10, 181, 251))
        self.label_2.setStyleSheet("")
        self.label_2.setText("")
        self.label_2.setPixmap(QtGui.QPixmap())
        self.label_2.setScaledContents(True)
        self.label_2.setObjectName("label_2")
        self.frame = QtWidgets.QFrame(Analyse)
        self.frame.setGeometry(QtCore.QRect(80, 340, 421, 231))
        self.frame.setStyleSheet("\n"
"background-color: rgba(202, 255, 187,150);")
        self.frame.setFrameShape(QtWidgets.QFrame.StyledPanel)
        self.frame.setFrameShadow(QtWidgets.QFrame.Raised)
        self.frame.setObjectName("frame")
        self.pushButton = QtWidgets.QPushButton(self.frame)
        self.pushButton.setGeometry(QtCore.QRect(160, 60, 93, 28))
        font = QtGui.QFont()
        font.setFamily("Monotype Corsiva")
        font.setPointSize(12)
        font.setBold(False)
        font.setItalic(True)
        font.setWeight(50)
        self.pushButton.setFont(font)
        self.pushButton.setStyleSheet("background-color: rgba(42, 42, 255,150);\n"
"color: rgb(217, 217, 217);\n"
"")
        self.pushButton.setObjectName("pushButton")
        self.pushButton.clicked.connect(self.selectImage)
        self.label_3 = QtWidgets.QLabel(self.frame)
        self.label_3.setGeometry(QtCore.QRect(100, 10, 221, 51))
        self.label_3.setStyleSheet("\n"
"background-color: rgba(255, 255, 255, 0);")
        self.label_3.setObjectName("label_3")
        self.label_4 = QtWidgets.QLabel(self.frame)
        self.label_4.setGeometry(QtCore.QRect(10, 90, 401, 51))
        self.label_4.setStyleSheet("\n"
"background-color: rgba(255, 255, 255, 0);")
        self.label_4.setWordWrap(True)
        self.label_4.setObjectName("label_4")
        self.pushButton_2 = QtWidgets.QPushButton(self.frame)
        self.pushButton_2.setGeometry(QtCore.QRect(130, 160, 161, 28))
        font = QtGui.QFont()
        font.setFamily("Monotype Corsiva")
        font.setPointSize(12)
        font.setBold(False)
        font.setItalic(True)
        font.setWeight(50)
        self.pushButton_2.setFont(font)
        self.pushButton_2.setStyleSheet("background-color: rgba(42, 42, 255,150);\n"
"color: rgb(217, 217, 217);\n"
"")
        self.pushButton_2.setObjectName("pushButton_2")
        self.pushButton_2.clicked.connect(self.analyse)
        self.label_5 = QtWidgets.QLabel(self.frame)
        self.label_5.setGeometry(QtCore.QRect(10, 200, 411, 20))
        font = QtGui.QFont()
        font.setFamily("Rockwell")
        font.setItalic(True)
        self.label_5.setFont(font)
        self.label_5.setStyleSheet("\n"
"background-color: rgba(255, 255, 255, 0);")
        self.label_5.setAlignment(QtCore.Qt.AlignCenter)
        self.label_5.setObjectName("label_5")
        self.pushButton_4 = QtWidgets.QPushButton(Analyse)
        self.pushButton_4.setGeometry(QtCore.QRect(240, 600, 93, 28))
        font = QtGui.QFont()
        font.setFamily("Monotype Corsiva")
        font.setPointSize(12)
        font.setBold(False)
        font.setItalic(True)
        font.setWeight(50)
        self.pushButton_4.setFont(font)
        self.pushButton_4.setStyleSheet("background-color: rgba(42, 42, 255,150);\n"
"color: rgb(217, 217, 217);\n"
"\n"
"")
        self.pushButton_4.setObjectName("pushButton_4")
        self.pushButton_4.clicked.connect(self.backhome)

        self.retranslateUi(Analyse)
        QtCore.QMetaObject.connectSlotsByName(Analyse)
    def backhome(self):
        Ui_Analyse.home=1
    def analyse(self):
        self.classify()
        
        
        
    def selectImage(self):
        print("Input Image")
        self.result=''
        self.fname = QFileDialog.getOpenFileName(None, 'Open file', '.','Image files (*.jpg *.jpeg)')
        self.label_2.setPixmap(QtGui.QPixmap(self.fname[0]))
        if(len(self.fname[0])>0):
            self.imagestatus="Image Selected"
        else:
            self.imagestatus="Select Scanned Image"
        self.gif2=QMovie("Gui-Images/doctor_checkup.gif")
        self.label.setMovie(self.gif2)
        self.gif2.start()
        self.label_3.setText(self._translate("Analyse", "<html><head/><body><p align=\"center\"><span style=\" font-size:11pt; color:#252a55;\">"+self.imagestatus+"</span></p></body></html>"))
        self.label_4.setText(self._translate("Analyse", "<html><head/><body><p align=\"center\"><span style=\" font-size:11pt; color:#252a55;\">"+self.result+"</span></p></body></html>"))
        self.label_5.setText(self._translate("Analyse", "<html><head/><body><p><span style=\" color:#005500;\">"+self.mailnotify+"</span></p></body></html>"))
    
    def classify(self):
        cnt=0
        if not(self.imagestatus=="Image Selected"):
            QMessageBox.information(None,'Warning!!','Image Not Selected , Select the image')
        else:
            image_path = self.fname[0]
        
        # Read the image_data

            image_data = tf.io.gfile.GFile(image_path, 'rb').read()
            
            
            # Loads label file, strips off carriage return
            label_lines = [line.rstrip() for line
                               in tf.io.gfile.GFile("logs/output_labels.txt")]
            
            # Unpersists graph from file
            with tf.io.gfile.GFile("logs/output_graph.pb", 'rb') as f:
                graph_def =  tf.compat.v1.GraphDef()
                graph_def.ParseFromString(f.read())
                _ = tf.import_graph_def(graph_def, name='')
            
            with tf.Session() as sess:
                # Feed the image_data as input to the graph and get first prediction
                softmax_tensor = sess.graph.get_tensor_by_name('final_result:0')
            
                predictions = sess.run(softmax_tensor, \
                         {'DecodeJpeg/contents:0': image_data})
            
                # Sort to show labels of first prediction in order of confidence
                top_k = predictions[0].argsort()[-len(predictions[0]):][::-1]
            
                for node_id in top_k:
                    human_string = label_lines[node_id]
                    score = predictions[0][node_id]
                    print('%s (score = %.5f)' % (human_string, score))
                    print("score is ",score)
                    if cnt==0:
                        self.predict=human_string
                        cnt=cnt+1
                self.label_2.setPixmap(QtGui.QPixmap(''))
                print("Prediction:"+self.predict)
                res=self.predict
                
                if(self.predict !='no tumor'):
                    self.result="Predicted Disease:"+self.predict[:].upper()
                   # self.doctor=self.experts[self.predict[:-1]][0]
                    #self.doctormail=self.experts[self.predict[:-1]][1]
                    self.makeReport()
                    self.mailnotify="Mail Sent"
                    suggest.main(res,self.mailid)
                    
                else:
                    self.result="No Worries ..! Patient is fit"
                print(self.patient,self.age,self.number,self.mailid)
                
                #self.retranslateUi(Analyse)
                self.gif2=QMovie("Gui-Images/analysis_complete.gif")
                self.label.setMovie(self.gif2)
                self.gif2.start()

                self.label_3.setText(self._translate("Analyse", "<html><head/><body><p align=\"center\"><span style=\" font-size:11pt; color:#252a55;\">"+self.imagestatus+"</span></p></body></html>"))
                self.label_4.setText(self._translate("Analyse", "<html><head/><body><p align=\"center\"><span style=\" font-size:11pt; color:#252a55;\">"+self.result+"</span></p></body></html>"))
                self.label_5.setText(self._translate("Analyse", "<html><head/><body><p><span style=\" color:#005500;\">"+self.mailnotify+"</span></p></body></html>"))
                self.pushButton_2.setText(self._translate("Analyse", self.buttonText)) 
                suggest.store(res,self.mailid)
    def makeReport(self):
        document = Document()
        section = document.sections[0]
        header = section.header
        paragraph = header.paragraphs[0]
        paragraph.text = "\BBHC"
        
        document.add_heading('\t\t\tSystem Generated Report', 0)
        
        p1 = document.add_paragraph('This is a automatic system generated ',style='List Bullet')
        p1.add_run('Report').bold = True
        p2 = document.add_paragraph('.Here the Analysis is made based on ',style='List Bullet')
        p2.add_run('CNN Algorithm.').italic = True
        
        document.add_heading('Patient Details-', level=1)
        
        
        document.add_paragraph('\tName:'+self.patient, style='Intense Quote')
        document.add_paragraph('\tAge:'+self.age, style='Intense Quote')
        document.add_paragraph('\tContact details:'+self.number, style='Intense Quote')
        document.add_paragraph('\tEmail ID:'+self.mailid, style='Intense Quote')
        
        document.add_heading('Scanned Image-', level=1)
        document.add_picture(self.fname[0], width=Inches(3),height=Inches(2))
          
        document.add_heading('Result-', level=1)
        p1 = document.add_paragraph(style='List Bullet')
        p1.add_run('Infection/Disease detected: ').italic=True
        #p1.add_run(self.predict[:-1])
        p1.add_run(self.predict[:].upper()).bold = True
        #d1 = document.add_paragraph(style='List Bullet')
        if(self.predict[:-1]=='glioma tumo'):
           d1 = document.add_paragraph(style='List Bullet')
           d1.add_run('Symptoms: ').bold=True
           d1.add_run('Headache\nNausea or vomitin\nConfusion or a decline in brain function\nMemory loss\nPersonality changes or irritability\nDifficulty with balance\nUrinary incontinence\nVision problems, such as blurred vision, double vision or loss of peripheral vision.')
           d2 = document.add_paragraph(style='List Bullet')
           d2.add_run('Causes:').bold=True
           d2.add_run('Gliomas are caused by the accumulation of genetic mutations in glial stem or progenitor cells, leading to their uncontrolled growth. Mutated genes are typically involved in functions such as tumor suppression, DNA repair, and regulation of cell growth.')
           d3 = document.add_paragraph(style='List Bullet')
           d3.add_run('Best treatment for glioma: ').bold=True
           d3.add_run('Chemotherapy, including GLIADEL® wafers and targeted therapy, is recommended for some high-grade gliomas after surgery and radiation therapy. Three types of chemotherapy may be used to treat glioma: Systemic, or standard, chemotherapy.\n')
           d4 = document.add_paragraph(style='List Bullet')
           d4.add_run('Types of Surgery for Gliomas: ').bold=True
           d4.add_run('Craniotomy\nIntraoperative brain mapping (awake brain surgery)')
           d5 = document.add_paragraph(style='List Bullet')
           d5.add_run('Radiation Therapy for Gliomas').bold=True
           d5.add_run('Internal radiation with the GliaSite Radiation Therapy System\nExternal beam radiation therapy\nStereotactic radiosurgery')
           d6 = document.add_paragraph(style='List Bullet')
           d6.add_run('Chemotherapy for Gliomas: ').bold=True
           d6.add_run('Systemic, or standard, chemotherapy\nGLIADEL® wafers\nTargeted therapy')
        elif(self.predict[:-1]=='meningioma tumo'):
            a1 = document.add_paragraph(style='List Bullet')
            a1.add_run('Symptoms: ').bold=True
            a1.add_run('Changes in vision, such as seeing double or blurriness\nHeadaches, especially those that are worse in the morning\nHearing loss or ringing in the ears\nMemory loss\nLoss of smell\nSeizures\nWeakness in your arms or legs.')
            a2 = document.add_paragraph(style='List Bullet')
            a2.add_run('Causes:').bold=True
            a2.add_run('The cause of meningiomas is not known. Exposure to radiation, especially in childhood, is the only known environmental risk factor for developing meningiomas. People who have a genetic condition, called neurofibromatosis type 2, are at increased risk for developing meningiomas.')
            a3 = document.add_paragraph(style='List Bullet')
            a3.add_run('Best treatment for meningioma tumor').bold=True
            a3.add_run('Conventional radiation therapy\nIntensity modulated radiation therapy (IMRT)\n3-dimensional conformal radiation therapy\nStereotactic radiosurgery\nFractionated stereotactic radiation therapy\nProton radiation therapy')
        elif(self.predict[:-1]=='pituitary tumo'):
           b1 = document.add_paragraph(style='List Bullet')
           b1.add_run('Symptoms: ').bold=True
           b1.add_run('Nausea and vomiting\nWeakness\nFeeling cold\nLess frequent or no menstrual periods\nSexual dysfunction\nIncreased amount of urine\nUnintended weight loss or gain.')
           b2 = document.add_paragraph(style='List Bullet')
           b2.add_run('Causes: ').bold=True
           b2.add_run('The causes of pituitary tumors are unknown. Some tumors are caused by hereditary disorders such as multiple endocrine neoplasia I (MEN I). The pituitary gland can be affected by other brain tumors that develop in the same part of the brain (skull base), resulting in similar symptoms.')
           b3 = document.add_paragraph(style='List Bullet')
           b3.add_run('Best treatment for pituitary: ').bold=True
           b3.add_run('Doctors generally use surgery, radiation therapy and medications, either alone or in combination, to treat a pituitary tumor and return hormone production to normal levels.')
           b4 = document.add_paragraph(style='List Bullet')
           b4.add_run('Types of Surgery for pituitary tumor:').bold=True
           b4.add_run('Endoscopic transnasal transsphenoidal approach\nTranscranial approach (craniotomy)')
           b5 = document.add_paragraph(style='List Bullet')
           b5.add_run('Types of Radiation therapy:').bold=True
           b5.add_run('Stereotactic radiosurgery\nExternal beam radiation\nIntensity modulated radiation therapy (IMRT)\nProton beam therapy')
           # d1.add_run('Specilist Assigned: ').italic=True
#        d1.add_run(self.doctor).bold = True
        footer_section = document.sections[0]
        footer = footer_section.footer
        footer_text = footer.paragraphs[0]
        footer_text.text = "\tNote:It is an Automatic Generated Report "
        
        document.save("Reports/"+self.patient+'_report.docx')
        self.sendmail("Reports/"+self.patient+'_report.docx')
        #self.sendmail("Precautions/"+self.predict[:-1]+'.docx',"P")
    
    def sendmail(self,filename):
        fromaddr = 'babymonitoringp@gmail.com'     #https://www.google.com/settings/security/lesssecureapps
        
        subject="System Generated Report"
        body = "Result of Patient Named "+self.patient+" Seems to be suffering from "+self.predict[:]+"\n Attachment:System Generated Report"
        # instance of MIMEMultipart 
        msg = MIMEMultipart() 
        # storing the senders email address   
        msg['From'] = "Hospital Update(Project)" 
          
        # storing the receivers email address  
        msg['To'] = self.mailid 
          
        # storing the subject  
        msg['Subject'] = subject
        # string to store the body of the mail
        
          
        # attach the body with the msg instance 
        msg.attach(MIMEText(body, 'plain')) 
          
        # open the file to be sent  
        attachment = open(filename, "rb") 
          
        # instance of MIMEBase and named as p 
        p = MIMEBase('application', 'octet-stream') 
          
        # To change the payload into encoded form 
        p.set_payload((attachment).read()) 
          
        # encode into base64 
        encoders.encode_base64(p) 
             
        # attach the instance 'p' to instance 'msg'
        
        p.add_header('Content-Disposition', "attachment; filename= %s" % filename)
        msg.attach(p) 
          
        # creates SMTP session 
        s = smtplib.SMTP('smtp.gmail.com', 587) 
          
        # start TLS for security 
        s.starttls() 
          
        # Authentication 
        s.login(fromaddr,"ryfyinhnxzpckcmw")#"Raspberry@123" 
          
        # Converts the Multipart msg into a string 
        text = msg.as_string() 
          
        # sending the mail 
        s.sendmail(fromaddr, self.mailid, text) 
          
        # terminating the session 
        s.quit()
        
    def retranslateUi(self, Analyse):
        self._translate = QtCore.QCoreApplication.translate
        Analyse.setWindowTitle(self._translate("Analyse", "E Report"))
        self.pushButton.setText(self._translate("Analyse", "Select"))
        self.label_3.setText(self._translate("Analyse", "<html><head/><body><p align=\"center\"><span style=\" font-size:11pt; color:#252a55;\">"+self.imagestatus+"</span></p></body></html>"))
        self.label_4.setText(self._translate("Analyse", "<html><head/><body><p align=\"center\"><span style=\" font-size:11pt; color:#252a55;\">"+self.result+"</span></p></body></html>"))
        self.pushButton_2.setText(self._translate("Analyse", self.buttonText))
        self.label_5.setText(self._translate("Analyse", "<html><head/><body><p><span style=\" color:#005500;\">"+self.mailnotify+"</span></p></body></html>"))
        self.pushButton_4.setText(self._translate("Analyse", "Home Page"))


if __name__ == "__main__":
    import sys
    app = QtWidgets.QApplication(sys.argv)
    Analyse = QtWidgets.QWidget()
    ui = Ui_Analyse('Name','Age','Number','pavanshettyasodu007007@gmail.com')
    ui.setupUi(Analyse)
    Analyse.show()
    sys.exit(app.exec_())

